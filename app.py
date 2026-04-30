# app.py completo con:
# - Login conservado
# - PIN corregido con SOPORTE_PASSWORD
# - Dashboard con diseño administrativo clásico
# - Pantalla Estudiantes y QR con diseño tipo registro de alumnos

from datetime import datetime, timedelta
from io import BytesIO
import os
import random
import smtplib
from email.message import EmailMessage
from zoneinfo import ZoneInfo

import qrcode
from docx import Document
from flask import Flask, request, redirect, session, send_file
from flask_sqlalchemy import SQLAlchemy
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from sqlalchemy import func, text


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "edutrack_secret_2026")

DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///edutrack.db")
if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

BOGOTA = ZoneInfo("America/Bogota")

APP_NAME = "EduTrack QR"
DESARROLLADOR = "Sebastián López / StudyTask"
SLOGAN = "Tecnología para una educación moderna y responsable"

INST_NOMBRE = "INSTITUCIÓN EDUCATIVA NOVA FUTURO"
INST_SEDE = "PRINCIPAL"
INST_DANE = "199999999999"
INST_NIT = "900.999.999-1"
INST_DIRECCION = "Calle 127 # 45-67, Bogotá D.C., Colombia"
INST_RESOLUCION = "Resolución institucional demo 2026"

SOPORTE_EMAIL = os.getenv("SOPORTE_EMAIL", "studytasksoporte@gmail.com")
SOPORTE_PASSWORD = os.getenv("SOPORTE_PASSWORD", "")

CARTERA_EMAIL = "carteraedutrackqr@gmail.com"
CARTERA_TELEFONO = "3126285480"
SOPORTE_TELEFONO = "3105615621"


CSS = """
<style>
:root{
--azul:#062b63;
--azul2:#0b63ce;
--azul3:#eaf2ff;
--rojo:#dc2626;
--verde:#16a34a;
--amarillo:#fbbf24;
--n:#111827;
--bg:#f5f7fb;
--s:0 18px 45px rgba(15,23,42,.12);
}
*{box-sizing:border-box}
body{margin:0;font-family:Segoe UI,Arial,sans-serif;background:var(--bg);color:var(--n)}
a{color:var(--azul2);font-weight:800}
input,select,textarea{width:100%;padding:13px;border:1px solid #d1d5db;border-radius:12px;margin:8px 0;font-size:15px;background:white}
button,.btn{background:var(--azul);color:white;border:0;border-radius:12px;padding:13px 17px;font-weight:800;text-decoration:none;display:inline-block;cursor:pointer}
.btn-red{background:var(--rojo)!important}
.btn-green{background:var(--verde)!important}
.center{display:flex;align-items:center;justify-content:center;min-height:100vh;padding:20px;background:linear-gradient(135deg,#061b3f,#0b63ce)}
.card{background:white;border-radius:24px;padding:28px;box-shadow:var(--s)}
.login-card{width:100%;max-width:440px}
.logo{width:92px;height:92px;object-fit:contain;background:white;border-radius:20px;padding:8px}
.hero-logo{background:linear-gradient(135deg,var(--azul),var(--azul2));color:white;border-radius:22px;padding:25px;text-align:center;border-bottom:6px solid var(--amarillo)}
.hero-logo h2,.hero-logo p{color:white}
.error{color:var(--rojo);font-weight:800}
.msg{padding:13px;border-radius:13px;font-weight:800;margin:12px 0}
.ok{background:#dcfce7;color:#047857}.warn{background:#fef3c7;color:#b45309}.danger{background:#fee2e2;color:#b91c1c}

.layout{display:flex;min-height:100vh;background:#d9d9d9}
.sidebar{width:230px;background:#08224b;color:white;padding:18px;border-right:5px solid #f5b400;position:sticky;top:0;height:100vh;overflow-y:auto}
.sidebar img{width:95px;height:95px;object-fit:contain;background:white;padding:4px;margin-bottom:12px}
.sidebar h2{font-size:24px;margin:4px 0}.sidebar p{color:#dbeafe;margin-top:0;font-size:13px}
.sidebar a{display:block;color:white;text-decoration:none;padding:12px 13px;border-radius:9px;margin:8px 0;background:#1d3f70;font-size:14px}
.sidebar a:hover,.sidebar .active{background:#1267d8}.brand-box{background:rgba(255,255,255,.09);border-radius:10px;padding:12px;margin-top:18px;font-size:13px}
.main{flex:1;padding:18px;overflow:auto}

.header{background:white;border-radius:18px;padding:18px 22px;box-shadow:var(--s);border-bottom:4px solid var(--amarillo);display:grid;grid-template-columns:1fr auto;gap:18px;margin-bottom:18px}
.header-info{display:flex;gap:16px;align-items:center}.header-info img{width:72px;height:72px;object-fit:contain}.header h1{margin:0;color:var(--azul);font-size:25px}.header p{margin:4px 0;color:#334155;font-size:13px}.user-badge{background:#eff6ff;color:var(--azul);padding:9px 13px;border-radius:999px;font-weight:800;display:inline-block;margin-bottom:10px;font-size:13px}

/* DISEÑO ADMINISTRATIVO CLÁSICO */
.classic-window{background:#bfbfbf;border:1px solid #555;border-radius:5px;box-shadow:inset 0 0 0 1px #e9e9e9,0 10px 28px rgba(0,0,0,.18);padding:6px;margin-bottom:16px;color:#111;font-family:Arial,sans-serif}
.classic-title{font-size:13px;background:linear-gradient(#f5f5f5,#b8b8b8);border:1px solid #8a8a8a;padding:6px 8px;margin-bottom:4px}
.classic-toolbar{height:31px;background:linear-gradient(#ececec,#c6c6c6);border:1px solid #9b9b9b;display:flex;align-items:center;gap:8px;padding:3px 8px;margin-bottom:10px}
.tool-btn{width:23px;height:23px;border:1px solid #9d9d9d;background:#e5e5e5;border-radius:2px;text-align:center;line-height:21px;font-size:12px;color:#333}
.record-box{display:inline-flex;align-items:center;gap:5px}.record-box input{width:55px;padding:4px;border-radius:0;margin:0;height:24px;font-size:12px;background:#e6e6e6}
.fieldset{border:1px solid #eee;margin:12px 18px 18px;padding:14px 18px 16px;position:relative}
.fieldset legend{font-size:12px;padding:0 8px;color:#111}.classic-grid{display:grid;grid-template-columns:120px 1fr 130px 1fr;gap:7px 10px;align-items:center}.classic-grid label{font-size:12px;text-align:right}.classic-grid input,.classic-grid select,.classic-grid textarea{border:1px solid #aaa;border-radius:0;background:#f4f4f4;height:23px;padding:3px 6px;margin:0;font-size:12px}.classic-grid textarea{height:45px;resize:vertical}.span3{grid-column:2/5}.span-all{grid-column:1/5}.classic-actions{padding:0 18px 15px;display:flex;gap:8px;flex-wrap:wrap}.classic-btn{background:#d8d8d8;color:#111;border:1px solid #777;border-radius:3px;padding:6px 12px;font-size:12px;font-weight:bold;text-decoration:none;display:inline-block}.classic-btn.primary{background:#0b63ce;color:white}.classic-btn.red{background:#cf1717;color:white}.classic-btn.green{background:#16a34a;color:white}.classic-btn.yellow{background:#ffe28a;color:#111}.classic-info{background:#eeeeee;border:1px solid #aaa;padding:8px;font-size:12px}
.dashboard-grid{display:grid;grid-template-columns:1fr 1fr;gap:14px}.classic-stat-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:8px}.classic-stat{border:1px solid #777;background:#eee;padding:10px;text-align:center}.classic-stat h3{font-size:28px;margin:0;color:#111}.classic-stat p{margin:3px 0 0;font-size:12px}.classic-stat.green{background:#d9f99d}.classic-stat.yellow{background:#fef08a}.classic-stat.red{background:#fecaca}.classic-stat.blue{background:#bfdbfe}
.finance-strip{border:1px solid #333;margin:10px 0;background:#d4d4d4;font-size:14px}.finance-row{display:grid;grid-template-columns:1.7fr 1fr;border-bottom:1px solid #333}.finance-row:last-child{border-bottom:0}.finance-row div{padding:5px;border-right:1px solid #333}.finance-row div:last-child{border-right:0;text-align:right;font-weight:bold}.finance-red{background:#c90000;color:white;font-weight:bold}.finance-yellow{background:#ffd21f;font-weight:bold}.finance-soft{background:#fff1a8}
.table-card{background:white;border-radius:12px;padding:16px;box-shadow:var(--s);margin-bottom:16px}.table-card h2{color:var(--azul);margin-top:0}table{width:100%;border-collapse:collapse;margin-top:10px}th{background:var(--azul);color:white;padding:10px;font-size:12px}td{padding:8px;border-bottom:1px solid #e5e7eb;text-align:center;font-size:13px}.estado{padding:5px 9px;border-radius:999px;font-weight:800;font-size:12px}.estado-temprano{background:#dcfce7;color:#047857}.estado-tarde{background:#fef3c7;color:#b45309}.estado-no{background:#fee2e2;color:#b91c1c}.qr-img{width:76px}.danger-link{color:var(--rojo)}
.footer{margin-top:20px;text-align:center;color:#64748b;font-size:13px}.footer strong{color:var(--azul)}.portal{width:100%;max-width:460px;text-align:center}.portal #reader{max-width:330px;margin:16px auto;border-radius:18px;overflow:hidden}.carnet{width:360px;background:white;border-radius:28px;padding:24px;text-align:center;box-shadow:var(--s);border-top:8px solid var(--azul2)}.carnet-head{background:linear-gradient(135deg,var(--azul),var(--azul2));color:white;border-radius:22px;padding:18px;border-bottom:6px solid var(--amarillo)}.carnet .qr{width:180px;margin:18px auto}.print-wrap{display:flex;align-items:center;justify-content:center;min-height:100vh;background:#eaf2ff;flex-direction:column}.contact-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:20px}.contact-card{background:white;border-radius:20px;padding:22px;box-shadow:var(--s);border-top:5px solid var(--azul2)}.contact-icon{width:50px;height:50px;border-radius:15px;background:var(--azul2);color:white;display:flex;align-items:center;justify-content:center;font-size:22px;margin-bottom:12px}.slogan-box{margin-top:25px;background:linear-gradient(135deg,var(--azul),var(--azul2));color:white;border-radius:24px;padding:26px;text-align:center}
@media(max-width:1100px){.layout{flex-direction:column}.sidebar{width:100%;height:auto;position:relative;border-right:0;border-bottom:5px solid var(--amarillo)}.header,.dashboard-grid,.classic-grid,.contact-grid{grid-template-columns:1fr}.classic-grid label{text-align:left}.span3,.span-all{grid-column:auto}.classic-stat-grid{grid-template-columns:1fr 1fr}.main{padding:12px}}
@media print{.no-print{display:none}.print-wrap{background:white}.carnet{box-shadow:none}}
</style>
"""


class Estudiante(db.Model):
    __tablename__ = "estudiantes"
    id = db.Column(db.Integer, primary_key=True)
    codigo = db.Column(db.String(80), unique=True, nullable=False)
    nombre = db.Column(db.String(120), nullable=False)
    apellido = db.Column(db.String(120), nullable=False)
    grado = db.Column(db.String(30), nullable=False)
    director = db.Column(db.String(120), nullable=False)


class Usuario(db.Model):
    __tablename__ = "usuarios"
    id = db.Column(db.Integer, primary_key=True)
    usuario = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    rol = db.Column(db.String(60), nullable=False)
    correo = db.Column(db.String(160), default="")
    grupo_docente = db.Column(db.String(30), default="")


class IngresoPorteria(db.Model):
    __tablename__ = "ingresos"
    id = db.Column(db.Integer, primary_key=True)
    estudiante_id = db.Column(db.Integer, db.ForeignKey("estudiantes.id"), nullable=False)
    fecha = db.Column(db.String(20), nullable=False)
    hora = db.Column(db.String(20), nullable=False)
    dia = db.Column(db.String(30), nullable=False)
    estado = db.Column(db.String(30), nullable=False)
    periodo = db.Column(db.String(30), nullable=False)
    registrado_por = db.Column(db.String(120), default="Portal móvil")
    estudiante = db.relationship("Estudiante")


class AsistenciaClase(db.Model):
    __tablename__ = "asistencias_clase"
    id = db.Column(db.Integer, primary_key=True)
    estudiante_id = db.Column(db.Integer, db.ForeignKey("estudiantes.id"), nullable=False)
    docente = db.Column(db.String(120), nullable=False)
    grupo = db.Column(db.String(30), nullable=False)
    fecha = db.Column(db.String(20), nullable=False)
    hora = db.Column(db.String(20), nullable=False)
    estado = db.Column(db.String(30), nullable=False)
    periodo = db.Column(db.String(30), nullable=False)
    observacion = db.Column(db.Text, default="")
    estudiante = db.relationship("Estudiante")


class Excusa(db.Model):
    __tablename__ = "excusas"
    id = db.Column(db.Integer, primary_key=True)
    estudiante_id = db.Column(db.Integer, db.ForeignKey("estudiantes.id"), nullable=False)
    fecha = db.Column(db.String(20), nullable=False)
    motivo = db.Column(db.Text, nullable=False)
    registrado_por = db.Column(db.String(120), nullable=False)
    periodo = db.Column(db.String(30), nullable=False)
    estudiante = db.relationship("Estudiante")


class Recuperacion(db.Model):
    __tablename__ = "recuperaciones"
    id = db.Column(db.Integer, primary_key=True)
    usuario = db.Column(db.String(80), nullable=False)
    pin = db.Column(db.String(80), nullable=False)


class Configuracion(db.Model):
    __tablename__ = "configuracion"
    id = db.Column(db.Integer, primary_key=True)
    periodo_actual = db.Column(db.String(30), default="Periodo 2")
    jornada = db.Column(db.String(30), default="Mañana")


def page(title, body):
    return f"""<!doctype html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>{title}</title>{CSS}</head><body>{body}</body></html>"""


def footer():
    return f"""<div class="footer"><strong>{APP_NAME}</strong> © 2026 | {SLOGAN}<br>Desarrollado por <b>{DESARROLLADOR}</b></div>"""


def shell(content):
    return f"""
<div class="layout">
  <aside class="sidebar">
    <img src="/static/img/logo-colegio.png" alt="Logo">
    <h2>{APP_NAME}</h2>
    <p>Sistema Inteligente de Asistencia Escolar</p>
    <a href="/dashboard" class="active">Inicio</a>
    <a href="/portal" target="_blank">Portal móvil</a>
    <a href="/docente-login">Portal docente</a>
    <a href="/estudiantes">Estudiantes</a>
    <a href="/usuarios">Usuarios</a>
    <a href="/reportes">Reportes</a>
    <a href="/alertas">Alertas</a>
    <a href="/excusas">Excusas</a>
    <a href="/contacto">Contacto institucional</a>
    <a href="/soporte-login">Soporte técnico</a>
    <div class="brand-box"><h3>{APP_NAME}</h3><p>{SLOGAN}</p></div>
  </aside>
  <main class="main">{content}{footer()}</main>
</div>
"""


def ahora(): return datetime.now(BOGOTA)
def fecha_hoy(): return ahora().strftime("%Y-%m-%d")
def hora_actual(): return ahora().strftime("%H:%M:%S")
def fecha_linda(): return ahora().strftime("%d/%m/%Y")


def migrar_columnas():
    try:
        if "postgresql" in DATABASE_URL:
            comandos = [
                "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS correo VARCHAR(160) DEFAULT ''",
                "ALTER TABLE usuarios ADD COLUMN IF NOT EXISTS grupo_docente VARCHAR(30) DEFAULT ''",
                "ALTER TABLE ingresos ADD COLUMN IF NOT EXISTS registrado_por VARCHAR(120) DEFAULT 'Portal móvil'",
                "ALTER TABLE configuracion ADD COLUMN IF NOT EXISTS periodo_actual VARCHAR(30) DEFAULT 'Periodo 2'",
                "ALTER TABLE configuracion ADD COLUMN IF NOT EXISTS jornada VARCHAR(30) DEFAULT 'Mañana'"
            ]
            for c in comandos:
                try: db.session.execute(text(c))
                except Exception: db.session.rollback()
            db.session.commit()
    except Exception as e:
        print("MIGRACION:", e); db.session.rollback()


def inicializar_bd():
    db.create_all(); migrar_columnas()
    if not Configuracion.query.first(): db.session.add(Configuracion(periodo_actual="Periodo 2", jornada="Mañana"))
    base = [("admin","1234","Rectoría","studytasksoporte@gmail.com",""),("soporte","1234","Soporte","studytasksoporte@gmail.com",""),("docente","1234","Docente","studytasksoporte@gmail.com","10")]
    for usuario,password,rol,correo,grupo in base:
        u = Usuario.query.filter(func.lower(Usuario.usuario) == usuario.lower()).first()
        if not u: db.session.add(Usuario(usuario=usuario,password=password,rol=rol,correo=correo,grupo_docente=grupo))
        else:
            if not u.correo: u.correo = correo
            if not u.grupo_docente and rol == "Docente": u.grupo_docente = grupo
    db.session.commit()


@app.before_request
def before(): inicializar_bd()


def config():
    c = Configuracion.query.first()
    if not c:
        c = Configuracion(periodo_actual="Periodo 2", jornada="Mañana")
        db.session.add(c); db.session.commit()
    return c


def periodo_actual(): return config().periodo_actual or "Periodo 2"
def jornada_actual(): return config().jornada or "Mañana"


def grados_disponibles():
    return sorted({e.grado for e in Estudiante.query.all()}, key=lambda x: int(str(x).split("-")[0]) if str(x).split("-")[0].isdigit() else 999)


def limpiar_codigo(texto):
    texto = (texto or "").strip()
    if "Codigo:" in texto: return texto.split("Codigo:")[1].split("\n")[0].strip()
    return texto


def qr_texto(e): return f"Codigo: {e.codigo}\nNombres: {e.nombre}\nApellidos: {e.apellido}\nGrado: {e.grado}\nDirector: {e.director}"


def login_usuario(usuario, password, rol_requerido=None):
    usuario=(usuario or "").strip(); password=(password or "").strip()
    u = Usuario.query.filter(func.lower(Usuario.usuario) == usuario.lower()).first()
    if not u or u.password.strip() != password: return None
    if rol_requerido and u.rol.lower().strip() != rol_requerido.lower().strip(): return None
    return u


def requiere_login(): return "usuario" in session
def rol_actual(): return session.get("rol", "")
def puede_admin(): return rol_actual() in ["Rectoría", "Administrador", "Soporte"]
def puede_estudiantes(): return rol_actual() in ["Rectoría", "Coordinación", "Secretaría", "Administrador", "Soporte"]


def estado_badge(estado):
    clase = "estado-temprano"
    if estado == "Tarde": clase = "estado-tarde"
    if estado == "No llegó": clase = "estado-no"
    return f'<span class="estado {clase}">{estado}</span>'


def registrar_ingreso(codigo, estado, registrado_por="Portal móvil"):
    codigo = limpiar_codigo(codigo)
    e = Estudiante.query.filter_by(codigo=codigo).first()
    if not e: return "Estudiante no registrado", "No registrado"
    db.session.add(IngresoPorteria(estudiante_id=e.id,fecha=fecha_hoy(),hora=hora_actual(),dia=ahora().strftime("%A"),estado=estado,periodo=periodo_actual(),registrado_por=registrado_por))
    db.session.commit()
    return f"Registro guardado: {e.nombre} {e.apellido} - {estado}", estado


def enviar_pin(correo_destino, pin):
    if not SOPORTE_EMAIL or not SOPORTE_PASSWORD: return False
    msg = EmailMessage(); msg["Subject"] = "PIN de recuperación - EduTrack QR"; msg["From"] = SOPORTE_EMAIL; msg["To"] = correo_destino
    msg.set_content(f"Tu PIN de recuperación es: {pin}\nEste PIN vence en 2 minutos.\n{APP_NAME}")
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(SOPORTE_EMAIL, SOPORTE_PASSWORD); smtp.send_message(msg)
        return True
    except Exception as e:
        print("ERROR PIN:", e); return False


@app.route("/")
def inicio(): return redirect("/login")


@app.route("/login", methods=["GET", "POST"])
def login():
    error = ""
    if request.method == "POST":
        user = login_usuario(request.form.get("usuario"), request.form.get("password"))
        if user:
            session["usuario"] = user.usuario; session["rol"] = user.rol; session["grupo_docente"] = user.grupo_docente or ""
            return redirect("/dashboard")
        error = "Usuario o contraseña incorrectos."
    return page("Login", f"""
<div class="center"><section class="card login-card"><div class="hero-logo"><img class="logo" src="/static/img/logo-colegio.png"><h2>{APP_NAME}</h2><p>{INST_NOMBRE}</p></div><h1>Iniciar sesión</h1><p>Acceso administrativo institucional.</p>{'<p class="error">'+error+'</p>' if error else ''}<form method="POST"><input name="usuario" placeholder="Usuario" required><input name="password" type="password" placeholder="Contraseña" required><button>Ingresar</button></form><p><a href="/recuperar">¿Olvidaste tu contraseña?</a> · <a href="/docente-login">Docentes</a> · <a href="/contacto">Contacto</a></p>{footer()}</section></div>
""")


@app.route("/recuperar", methods=["GET", "POST"])
def recuperar():
    mensaje = ""
    if request.method == "POST":
        usuario = request.form.get("usuario", "").strip(); correo = request.form.get("correo", "").strip()
        user = Usuario.query.filter(func.lower(Usuario.usuario) == usuario.lower(), Usuario.correo == correo).first()
        if user:
            pin = str(random.randint(100000, 999999)); vence = (ahora() + timedelta(minutes=2)).timestamp()
            Recuperacion.query.filter_by(usuario=user.usuario).delete(); db.session.add(Recuperacion(usuario=user.usuario, pin=f"{pin}|{vence}")); db.session.commit()
            if enviar_pin(correo, pin):
                session["recuperar_usuario"] = user.usuario; return redirect("/validar_pin")
            mensaje = "No se pudo enviar el correo. Revisa SOPORTE_EMAIL y SOPORTE_PASSWORD en Render."
        else: mensaje = "Usuario o correo no encontrado. Debe estar creado en Usuarios con ese correo."
    return page("Recuperar", f"""<div class="center"><section class="card login-card"><h1>Recuperar contraseña</h1><p>{mensaje}</p><form method="POST"><input name="usuario" placeholder="Usuario" required><input name="correo" placeholder="Correo registrado" required><button>Enviar PIN</button></form><p>El PIN vence en 2 minutos.</p><a href="/login">Volver</a></section></div>""")


@app.route("/validar_pin", methods=["GET", "POST"])
def validar_pin():
    if "recuperar_usuario" not in session: return redirect("/recuperar")
    mensaje = ""
    if request.method == "POST":
        r = Recuperacion.query.filter_by(usuario=session["recuperar_usuario"]).first()
        if r:
            partes = r.pin.split("|"); pin_real = partes[0]; vence = float(partes[1]) if len(partes) > 1 else 0
            if ahora().timestamp() > vence:
                Recuperacion.query.filter_by(usuario=session["recuperar_usuario"]).delete(); db.session.commit(); mensaje = "El PIN venció. Solicita uno nuevo."
            elif request.form.get("pin") == pin_real:
                u = Usuario.query.filter_by(usuario=session["recuperar_usuario"]).first(); u.password = request.form.get("password", "").strip()
                Recuperacion.query.filter_by(usuario=u.usuario).delete(); db.session.commit(); session.pop("recuperar_usuario", None); return redirect("/login")
            else: mensaje = "PIN incorrecto."
    return page("Validar PIN", f"""<div class="center"><section class="card login-card"><h1>Nueva contraseña</h1><p>{mensaje}</p><form method="POST"><input name="pin" placeholder="PIN recibido" required><input name="password" type="password" placeholder="Nueva contraseña" required><button>Guardar contraseña</button></form></section></div>""")


@app.route("/dashboard", methods=["GET", "POST"])
def dashboard():
    if not requiere_login(): return redirect("/login")
    cfg = config()
    if request.method == "POST":
        cfg.periodo_actual = request.form.get("periodo_actual", cfg.periodo_actual); cfg.jornada = request.form.get("jornada", cfg.jornada); db.session.commit(); return redirect("/dashboard")
    grupo = request.args.get("grupo", "TODOS"); periodo = request.args.get("periodo", periodo_actual()); hoy = fecha_hoy()
    estudiantes = Estudiante.query.all(); estudiantes_grupo = estudiantes if grupo == "TODOS" else [e for e in estudiantes if e.grado == grupo]
    ingresos_query = IngresoPorteria.query.filter_by(fecha=hoy)
    if periodo != "TODOS": ingresos_query = ingresos_query.filter_by(periodo=periodo)
    ingresos = ingresos_query.all()
    if grupo != "TODOS": ingresos = [i for i in ingresos if i.estudiante.grado == grupo]
    ids_ingreso = {i.estudiante_id for i in ingresos}
    total = len(estudiantes_grupo); temprano = sum(1 for i in ingresos if i.estado == "Temprano"); tarde = sum(1 for i in ingresos if i.estado == "Tarde"); no_llego = max(total - len(ids_ingreso), 0) + sum(1 for i in ingresos if i.estado == "No llegó")
    asistencias = AsistenciaClase.query.filter_by(fecha=hoy).all(); docentes = len({a.docente for a in asistencias}); grupos_asistencia = len({a.grupo for a in asistencias})
    opciones_grupo = '<option value="TODOS">TODOS</option>' + ''.join(f'<option value="{g}" {"selected" if grupo == g else ""}>{g}-01-MAÑANA</option>' for g in grados_disponibles())
    opciones_periodo = ''.join(f'<option value="{p}" {"selected" if periodo == p else ""}>{p}</option>' for p in ["TODOS", "Periodo 1", "Periodo 2", "Periodo 3"])
    filas_estudiantes = ''.join(f"""<tr><td>{e.codigo}</td><td>{e.nombre} {e.apellido}</td><td>{e.grado}</td><td>{e.director}</td><td>{estado_badge(next((i.estado for i in ingresos if i.estudiante_id == e.id), 'No llegó'))}</td><td>{next((i.hora for i in ingresos if i.estudiante_id == e.id), '-')}</td></tr>""" for e in estudiantes_grupo[:12])
    ultimos = IngresoPorteria.query.join(Estudiante).filter(IngresoPorteria.fecha == hoy).order_by(IngresoPorteria.hora.desc()).limit(9).all()
    filas_ultimos = ''.join(f"<tr><td>{i.estudiante.grado}</td><td>{i.estudiante.nombre} {i.estudiante.apellido}</td><td>{i.hora}</td><td>{estado_badge(i.estado)}</td></tr>" for i in ultimos)

    content = f"""
<div class="classic-window">
  <div class="classic-title">Panel Administrativo - {APP_NAME}</div>
  <div class="classic-toolbar"><span class="tool-btn">|&lt;</span><span class="tool-btn">&lt;</span><span class="record-box"><input value="{total}"><small>de {len(estudiantes)}</small></span><span class="tool-btn">&gt;</span><span class="tool-btn">&gt;|</span><span class="tool-btn">+</span><span class="tool-btn">✕</span><span class="tool-btn">💾</span><span class="tool-btn">✎</span><span class="tool-btn">●</span></div>
  <fieldset class="fieldset"><legend>Datos institucionales</legend>
    <div class="classic-grid">
      <label>Institución:</label><input value="{INST_NOMBRE}" readonly><label>Usuario:</label><input value="{session['usuario']} - {session['rol']}" readonly>
      <label>Sede:</label><input value="{INST_SEDE}" readonly><label>Fecha:</label><input value="{fecha_linda()} {ahora().strftime('%I:%M:%S %p')}" readonly>
      <label>DANE:</label><input value="{INST_DANE}" readonly><label>NIT:</label><input value="{INST_NIT}" readonly>
      <label>Dirección:</label><input class="span3" value="{INST_DIRECCION}" readonly>
    </div>
  </fieldset>
  <div class="dashboard-grid">
    <fieldset class="fieldset"><legend>Configuración manual del sistema</legend>
      <form method="POST"><div class="classic-grid"><label>Periodo:</label><select name="periodo_actual"><option {"selected" if periodo_actual()=="Periodo 1" else ""}>Periodo 1</option><option {"selected" if periodo_actual()=="Periodo 2" else ""}>Periodo 2</option><option {"selected" if periodo_actual()=="Periodo 3" else ""}>Periodo 3</option></select><label>Jornada:</label><select name="jornada"><option {"selected" if jornada_actual()=="Mañana" else ""}>Mañana</option><option {"selected" if jornada_actual()=="Tarde" else ""}>Tarde</option><option {"selected" if jornada_actual()=="Única" else ""}>Única</option></select></div><br><button class="classic-btn primary">Guardar configuración</button></form>
    </fieldset>
    <fieldset class="fieldset"><legend>Filtro de consulta</legend>
      <form method="GET"><div class="classic-grid"><label>Periodo:</label><select name="periodo">{opciones_periodo}</select><label>Grupo:</label><select name="grupo">{opciones_grupo}</select></div><br><button class="classic-btn primary">Consultar</button></form>
    </fieldset>
  </div>
  <fieldset class="fieldset"><legend>Resumen de asistencia</legend>
    <div class="classic-stat-grid"><div class="classic-stat blue"><h3>{total}</h3><p>Estudiantes</p></div><div class="classic-stat green"><h3>{temprano}</h3><p>Temprano</p></div><div class="classic-stat yellow"><h3>{tarde}</h3><p>Tarde</p></div><div class="classic-stat red"><h3>{no_llego}</h3><p>No llegó</p></div></div>
  </fieldset>
  <div class="finance-strip"><div class="finance-row"><div>Docentes con asistencia</div><div class="finance-yellow">{docentes}</div></div><div class="finance-row"><div>Grupos con asistencia</div><div class="finance-soft">{grupos_asistencia}</div></div><div class="finance-row"><div class="finance-red">Portal móvil de escaneo</div><div class="finance-red"><a style="color:white" href="/portal" target="_blank">ABRIR</a></div></div></div>
  <div class="classic-actions"><a class="classic-btn primary" href="/portal" target="_blank">Portal móvil</a><a class="classic-btn" href="/docente-login">Portal docente</a><a class="classic-btn green" href="/estudiantes">Crear estudiante o QR</a><a class="classic-btn" href="/reportes">Reportes</a><a class="classic-btn" href="/usuarios">Usuarios</a><a class="classic-btn red" href="/logout">Salir</a></div>
</div>
<div class="dashboard-grid"><section class="table-card"><h2>Estudiantes del grupo: {grupo}</h2><table><tr><th>Código</th><th>Nombre y apellido</th><th>Grado</th><th>Director</th><th>Estado</th><th>Hora</th></tr>{filas_estudiantes}</table></section><section class="table-card"><h2>Últimos ingresos</h2><table><tr><th>Grupo</th><th>Estudiante</th><th>Hora</th><th>Estado</th></tr>{filas_ultimos}</table></section></div>
"""
    return page("Dashboard", shell(content))


@app.route("/portal", methods=["GET", "POST"])
def portal():
    mensaje = ""; estado = ""
    if request.method == "POST": mensaje, estado = registrar_ingreso(request.form.get("codigo"), request.form.get("estado", "Temprano"), "Portal móvil")
    return page("Portal móvil", f"""<div class="center"><section class="card portal"><img class="logo" src="/static/img/logo-colegio.png"><h1>Portal de Ingreso</h1><p>Escanea el QR o escribe el código.</p><div id="reader"></div><form method="POST" id="portal-form"><input name="codigo" id="codigo" placeholder="Código del estudiante" required><select name="estado"><option>Temprano</option><option>Tarde</option><option>No llegó</option></select><button>Guardar ingreso</button></form>{'<div class="msg ok">'+mensaje+'</div>' if mensaje and estado!='No registrado' else ''}{'<div class="msg danger">'+mensaje+'</div>' if mensaje and estado=='No registrado' else ''}<p><a href="/login">Administración</a></p>{footer()}</section></div><script src="https://unpkg.com/html5-qrcode"></script><script>function onScanSuccess(decodedText){{document.getElementById("codigo").value=decodedText;document.getElementById("portal-form").submit();}}const scanner=new Html5Qrcode("reader");Html5Qrcode.getCameras().then(cameras=>{{if(cameras&&cameras.length){{scanner.start({{facingMode:"environment"}},{{fps:10,qrbox:250}},onScanSuccess);}}}}).catch(e=>alert("No se pudo abrir la cámara. Revisa permisos."));</script>""")


@app.route("/docente-login", methods=["GET", "POST"])
def docente_login():
    error = ""
    if request.method == "POST":
        u = login_usuario(request.form.get("usuario"), request.form.get("password"), "Docente")
        if u:
            session["usuario"] = u.usuario; session["rol"] = u.rol; session["grupo_docente"] = u.grupo_docente or ""; return redirect("/docente")
        error = "Usuario docente incorrecto."
    return page("Portal docente", f"""<div class="center"><section class="card login-card"><img class="logo" src="/static/img/logo-colegio.png"><h1>Portal Docente</h1><p>Asistencia en aula.</p>{'<p class="error">'+error+'</p>' if error else ''}<form method="POST"><input name="usuario" placeholder="Usuario docente" required><input name="password" type="password" placeholder="Contraseña" required><button>Ingresar</button></form><p>Prueba: docente / 1234</p><a href="/login">Volver</a>{footer()}</section></div>""")


@app.route("/docente", methods=["GET", "POST"])
def docente():
    if not requiere_login() or rol_actual() != "Docente": return redirect("/docente-login")
    grupos = grados_disponibles(); grupo = request.args.get("grupo") or session.get("grupo_docente") or (grupos[0] if grupos else ""); mensaje = ""
    if request.method == "POST":
        grupo = request.form.get("grupo")
        for e in Estudiante.query.filter_by(grado=grupo).all():
            estado = request.form.get(f"estado_{e.id}", "Presente"); obs = request.form.get(f"obs_{e.id}", "")
            db.session.add(AsistenciaClase(estudiante_id=e.id, docente=session["usuario"], grupo=grupo, fecha=fecha_hoy(), hora=hora_actual(), estado=estado, periodo=periodo_actual(), observacion=obs))
            if estado == "Excusa": db.session.add(Excusa(estudiante_id=e.id, fecha=fecha_hoy(), motivo=obs or "Excusa registrada", registrado_por=session["usuario"], periodo=periodo_actual()))
        db.session.commit(); mensaje = "Asistencia guardada correctamente."
    opciones = ''.join(f'<option value="{g}" {"selected" if g==grupo else ""}>{g}-01-MAÑANA</option>' for g in grupos)
    estudiantes = Estudiante.query.filter_by(grado=grupo).order_by(Estudiante.nombre.asc()).all()
    filas = ''.join(f"<tr><td>{e.codigo}</td><td>{e.nombre} {e.apellido}</td><td>{e.grado}</td><td>{e.director}</td><td><select name='estado_{e.id}'><option>Presente</option><option>Ausente</option><option>Tarde</option><option>Excusa</option></select></td><td><input name='obs_{e.id}' placeholder='Observación'></td></tr>" for e in estudiantes)
    content = f"""<header class="header"><div class="header-info"><img src="/static/img/logo-colegio.png"><div><h1>Portal Docente</h1><p>Docente: {session['usuario']} · Periodo: {periodo_actual()}</p></div></div><a class="btn btn-red" href="/logout">Salir</a></header>{'<div class="msg ok">'+mensaje+'</div>' if mensaje else ''}<div class="panel-box"><form method="GET"><label>Grupo</label><select name="grupo">{opciones}</select><button>Ver grupo</button></form></div><div class="table-card"><h2>Asistencia grupo {grupo}</h2><form method="POST"><input type="hidden" name="grupo" value="{grupo}"><table><tr><th>Código</th><th>Nombre</th><th>Grado</th><th>Director</th><th>Estado</th><th>Observación</th></tr>{filas}</table><br><button>Guardar asistencia</button></form></div>"""
    return page("Docente", shell(content))


@app.route("/estudiantes", methods=["GET", "POST"])
def estudiantes():
    if not requiere_login(): return redirect("/login")
    if not puede_estudiantes(): return "No tienes permiso."
    mensaje = ""
    if request.method == "POST":
        codigo = request.form.get("codigo", "").strip()
        if codigo:
            e = Estudiante.query.filter_by(codigo=codigo).first()
            if not e:
                e = Estudiante(codigo=codigo); db.session.add(e); mensaje = "Alumno registrado correctamente."
            else: mensaje = "Alumno actualizado correctamente."
            e.nombre = request.form.get("nombre", "").strip(); e.apellido = request.form.get("apellido", "").strip(); e.grado = request.form.get("grado", "").strip(); e.director = request.form.get("director", "").strip(); db.session.commit()
        return redirect("/estudiantes")
    total_est = Estudiante.query.count()
    filas = ''.join(f"""<tr><td>{e.codigo}</td><td>{e.nombre}</td><td>{e.apellido}</td><td>{e.grado}</td><td>{e.director}</td><td><img class="qr-img" src="/qr/{e.id}"><br><a href="/qr_descargar/{e.id}">Descargar</a></td><td><a href="/carnet/{e.id}">Carné</a> · <a href="/historial/{e.id}">Historial</a> · <a class="danger-link" href="/eliminar_estudiante/{e.id}">Eliminar</a></td></tr>""" for e in Estudiante.query.order_by(Estudiante.grado.asc(), Estudiante.nombre.asc()).all())
    content = f"""
<div class="classic-window">
  <div class="classic-title">Registro de Alumnos</div>
  <div class="classic-toolbar"><span class="tool-btn">|&lt;</span><span class="tool-btn">&lt;</span><span class="record-box"><input value="{total_est}"><small>de {total_est}</small></span><span class="tool-btn">&gt;</span><span class="tool-btn">&gt;|</span><span class="tool-btn">+</span><span class="tool-btn">✕</span><span class="tool-btn">💾</span><span class="tool-btn">✎</span><span class="tool-btn">🔴</span></div>
  {'<div class="msg ok">'+mensaje+'</div>' if mensaje else ''}
  <form method="POST">
    <fieldset class="fieldset"><legend>Datos Personales</legend><div class="classic-grid"><label>Código:</label><input name="codigo" required><label>Nombre:</label><input name="nombre" required><label>Apellido paterno:</label><input name="apellido" required><label>Apellido materno:</label><input placeholder="Opcional"><label>DNI Acudiente:</label><input class="span3" placeholder="Documento del acudiente"></div></fieldset>
    <fieldset class="fieldset"><legend>Datos académicos</legend><div class="classic-grid"><label>Fecha de ingreso:</label><input value="{fecha_hoy()}" readonly><label>Fecha de nacimiento:</label><input placeholder="__/__/____"><label>Colegio procedencia:</label><input placeholder="Colegio anterior"><label>Condición:</label><select><option>Activo</option><option>Retirado</option><option>Traslado</option></select><label>Observación:</label><input class="span3" placeholder="Observación"><label>Nivel y grado:</label><select name="grado" class="span3" required><option value="">Seleccione</option><option>6</option><option>7</option><option>8</option><option>9</option><option>10</option><option>11</option></select><label>Director:</label><input name="director" class="span3" required></div></fieldset>
    <fieldset class="fieldset"><legend>Datos de contacto</legend><div class="classic-grid"><label>Dirección:</label><input class="span3" placeholder="Dirección"><label>Correo:</label><input placeholder="Correo"><label>Teléfono:</label><input placeholder="Teléfono"></div></fieldset>
    <div class="classic-actions"><button class="classic-btn primary">Guardar alumno</button><a class="classic-btn green" href="/exportar_estudiantes">Exportar Excel</a><a class="classic-btn" href="/dashboard">Volver</a></div>
  </form>
</div>
<section class="table-card"><h2>Alumnos registrados y códigos QR</h2><table><tr><th>Código</th><th>Nombre</th><th>Apellido</th><th>Grado</th><th>Director</th><th>QR</th><th>Acciones</th></tr>{filas}</table></section>
"""
    return page("Estudiantes", shell(content))


@app.route("/eliminar_estudiante/<int:id>")
def eliminar_estudiante(id):
    if not requiere_login(): return redirect("/login")
    e = Estudiante.query.get_or_404(id); IngresoPorteria.query.filter_by(estudiante_id=e.id).delete(); AsistenciaClase.query.filter_by(estudiante_id=e.id).delete(); Excusa.query.filter_by(estudiante_id=e.id).delete(); db.session.delete(e); db.session.commit(); return redirect("/estudiantes")


@app.route("/qr/<int:id>")
def qr_estudiante(id):
    e = Estudiante.query.get_or_404(id); img = qrcode.make(qr_texto(e)); b = BytesIO(); img.save(b, format="PNG"); b.seek(0); return send_file(b, mimetype="image/png")


@app.route("/qr_descargar/<int:id>")
def qr_descargar(id):
    e = Estudiante.query.get_or_404(id); img = qrcode.make(qr_texto(e)); b = BytesIO(); img.save(b, format="PNG"); b.seek(0); return send_file(b, mimetype="image/png", as_attachment=True, download_name=f"{e.codigo}.png")


@app.route("/carnet/<int:id>")
def carnet(id):
    e = Estudiante.query.get_or_404(id)
    return page("Carné", f"""<div class="print-wrap"><div class="carnet"><div class="carnet-head"><img class="logo" src="/static/img/logo-colegio.png"><h3>{INST_NOMBRE}</h3><p>Sede {INST_SEDE}</p></div><h2>{e.nombre} {e.apellido}</h2><p><b>Grado:</b> {e.grado}</p><p><b>Director:</b> {e.director}</p><p><b>Código:</b> {e.codigo}</p><img class="qr" src="/qr/{e.id}"><p>{APP_NAME}</p></div><br><button class="no-print" onclick="window.print()">Imprimir carné</button><br><a class="no-print" href="/estudiantes">Volver</a></div>""")


@app.route("/usuarios", methods=["GET", "POST"])
def usuarios():
    if not requiere_login(): return redirect("/login")
    if not puede_admin(): return "No tienes permiso."
    if request.method == "POST":
        usuario = request.form.get("usuario", "").strip()
        if usuario and not Usuario.query.filter(func.lower(Usuario.usuario) == usuario.lower()).first():
            db.session.add(Usuario(usuario=usuario, password=request.form.get("password", "").strip(), rol=request.form.get("rol"), correo=request.form.get("correo", "").strip(), grupo_docente=request.form.get("grupo_docente", "").strip())); db.session.commit()
        return redirect("/usuarios")
    filas = ''.join(f"<tr><td>{u.usuario}</td><td>{u.rol}</td><td>{u.correo}</td><td>{u.grupo_docente}</td><td><a class='danger-link' href='/eliminar_usuario/{u.id}'>Eliminar</a></td></tr>" for u in Usuario.query.order_by(Usuario.rol.asc()).all())
    content = f"""<header class="header"><div class="header-info"><img src="/static/img/logo-colegio.png"><div><h1>Usuarios</h1><p>Administración de accesos.</p></div></div><a class="btn" href="/dashboard">Volver</a></header><div class="table-card"><h2>Crear usuario</h2><form method="POST"><input name="usuario" placeholder="Usuario" required><input name="password" type="password" placeholder="Contraseña" required><input name="correo" placeholder="Correo de recuperación"><select name="rol"><option>Rectoría</option><option>Coordinación</option><option>Secretaría</option><option>Docente</option><option>Soporte</option></select><input name="grupo_docente" placeholder="Grupo docente, ejemplo: 10"><button>Crear usuario</button></form></div><div class="table-card"><h2>Usuarios registrados</h2><table><tr><th>Usuario</th><th>Rol</th><th>Correo</th><th>Grupo docente</th><th>Acción</th></tr>{filas}</table></div>"""
    return page("Usuarios", shell(content))


@app.route("/eliminar_usuario/<int:id>")
def eliminar_usuario(id):
    if not requiere_login(): return redirect("/login")
    u = Usuario.query.get_or_404(id)
    if u.usuario != session.get("usuario"):
        db.session.delete(u); db.session.commit()
    return redirect("/usuarios")


@app.route("/historial/<int:id>")
def historial(id):
    if not requiere_login(): return redirect("/login")
    e = Estudiante.query.get_or_404(id); ingresos = IngresoPorteria.query.filter_by(estudiante_id=id).order_by(IngresoPorteria.fecha.desc()).all()
    filas = ''.join(f"<tr><td>{i.fecha}</td><td>{i.hora}</td><td>{i.estado}</td><td>{i.periodo}</td><td>{i.registrado_por}</td></tr>" for i in ingresos)
    return page("Historial", shell(f"<header class='header'><div class='header-info'><img src='/static/img/logo-colegio.png'><div><h1>Historial de {e.nombre} {e.apellido}</h1><p>Grado {e.grado}</p></div></div><a class='btn' href='/estudiantes'>Volver</a></header><div class='table-card'><h2>Ingresos</h2><table><tr><th>Fecha</th><th>Hora</th><th>Estado</th><th>Periodo</th><th>Quién reporta</th></tr>{filas}</table></div>"))


@app.route("/excusas")
def excusas():
    if not requiere_login(): return redirect("/login")
    filas = ''.join(f"<tr><td>{x.fecha}</td><td>{x.estudiante.codigo}</td><td>{x.estudiante.nombre} {x.estudiante.apellido}</td><td>{x.estudiante.grado}</td><td>{x.motivo}</td><td>{x.registrado_por}</td></tr>" for x in Excusa.query.order_by(Excusa.fecha.desc()).all())
    return page("Excusas", shell(f"<header class='header'><div class='header-info'><img src='/static/img/logo-colegio.png'><div><h1>Excusas</h1><p>Registro institucional.</p></div></div></header><div class='table-card'><table><tr><th>Fecha</th><th>Código</th><th>Nombre</th><th>Grado</th><th>Motivo</th><th>Registró</th></tr>{filas}</table></div>"))


@app.route("/alertas")
def alertas():
    if not requiere_login(): return redirect("/login")
    filas = ""
    for e in Estudiante.query.all():
        ingresos = IngresoPorteria.query.filter_by(estudiante_id=e.id, periodo=periodo_actual()).all(); aula = AsistenciaClase.query.filter_by(estudiante_id=e.id, periodo=periodo_actual()).all()
        tardes = sum(1 for i in ingresos if i.estado == "Tarde") + sum(1 for a in aula if a.estado == "Tarde"); ausencias = sum(1 for a in aula if a.estado == "Ausente")
        if tardes >= 3: filas += f"<tr><td>{e.codigo}</td><td>{e.nombre} {e.apellido}</td><td>{e.grado}</td><td>{tardes} llegadas tarde</td></tr>"
        if ausencias >= 3: filas += f"<tr><td>{e.codigo}</td><td>{e.nombre} {e.apellido}</td><td>{e.grado}</td><td>{ausencias} ausencias</td></tr>"
    return page("Alertas", shell(f"<header class='header'><div class='header-info'><img src='/static/img/logo-colegio.png'><div><h1>Alertas</h1><p>Novedades por asistencia.</p></div></div></header><div class='table-card'><table><tr><th>Código</th><th>Nombre</th><th>Grado</th><th>Alerta</th></tr>{filas}</table></div>"))


@app.route("/contacto")
def contacto():
    return page("Contacto", shell(f"""<header class="header"><div class="header-info"><img src="/static/img/logo-colegio.png"><div><h1>Contacto Institucional</h1><p>Soporte y acompañamiento técnico.</p></div></div></header><section class="card"><div class="contact-grid"><div class="contact-card"><div class="contact-icon">☎</div><h3>DESARROLLADOR PRINCIPAL</h3><p><b>Sebastián López</b><br>Soporte técnico EduTrack QR</p><p>📞 <a href="tel:{SOPORTE_TELEFONO}">{SOPORTE_TELEFONO}</a></p><p>✉️ <a href="mailto:{SOPORTE_EMAIL}">{SOPORTE_EMAIL}</a></p></div><div class="contact-card"><div class="contact-icon">✉</div><h3>SOPORTE TECNOLÓGICO</h3><p>Plataforma EduTrack QR</p><p>✉️ <a href="mailto:{SOPORTE_EMAIL}">{SOPORTE_EMAIL}</a></p></div><div class="contact-card"><div class="contact-icon">💼</div><h3>CARTERA Y FACTURACIÓN</h3><p>📞 <a href="tel:{CARTERA_TELEFONO}">{CARTERA_TELEFONO}</a></p><p>✉️ <a href="mailto:{CARTERA_EMAIL}">{CARTERA_EMAIL}</a></p></div></div><div class="slogan-box"><h2>{APP_NAME}</h2><p>{SLOGAN}</p></div></section>"""))


@app.route("/soporte-login", methods=["GET", "POST"])
def soporte_login():
    error = ""
    if request.method == "POST":
        u = login_usuario(request.form.get("usuario"), request.form.get("password"), "Soporte")
        if u:
            session["usuario"] = u.usuario; session["rol"] = u.rol; return redirect("/soporte")
        error = "Acceso incorrecto."
    return page("Soporte", f"<div class='center'><section class='card login-card'><h1>Soporte</h1><p class='error'>{error}</p><form method='POST'><input name='usuario' placeholder='Usuario'><input name='password' type='password' placeholder='Contraseña'><button>Ingresar</button></form><a href='/login'>Volver</a></section></div>")


@app.route("/soporte")
def soporte():
    if not requiere_login() or rol_actual() != "Soporte": return redirect("/soporte-login")
    return page("Soporte", shell(f"<header class='header'><div class='header-info'><img src='/static/img/logo-colegio.png'><div><h1>Panel de Soporte</h1><p>Estado técnico del sistema.</p></div></div></header><section class='table-card'><h2>Estado</h2><p>Usuarios: {Usuario.query.count()}</p><p>Estudiantes: {Estudiante.query.count()}</p><p>Ingresos: {IngresoPorteria.query.count()}</p></section>"))


def datos_reporte():
    data = []
    for i in IngresoPorteria.query.join(Estudiante).order_by(Estudiante.grado.asc(), IngresoPorteria.fecha.desc()).all():
        data.append([i.estudiante.codigo, f"{i.estudiante.nombre} {i.estudiante.apellido}", i.estudiante.grado, i.estudiante.director, i.fecha, i.hora, i.registrado_por])
    return data


@app.route("/reportes")
def reportes():
    if not requiere_login(): return redirect("/login")
    filas = ''.join(f"<tr><td>{r[0]}</td><td>{r[1]}</td><td>{r[2]}</td><td>{r[3]}</td><td>{r[4]}</td><td>{r[5]}</td><td>{r[6]}</td></tr>" for r in datos_reporte())
    content = f"<header class='header'><div class='header-info'><img src='/static/img/logo-colegio.png'><div><h1>Reportes institucionales</h1><p>Formato tipo planilla.</p></div></div><div><a class='btn' href='/exportar_excel_reportes'>Excel</a> <a class='btn' href='/exportar_pdf'>PDF</a> <a class='btn' href='/exportar_word'>Word</a></div></header><div class='table-card'><table><tr><th>Código</th><th>Nombre y apellido</th><th>Grado</th><th>Director</th><th>Fecha</th><th>Hora</th><th>Quién reporta</th></tr>{filas}</table></div>"
    return page("Reportes", shell(content))


def encabezado(): return [INST_NOMBRE, f"SEDE {INST_SEDE}", INST_RESOLUCION, f"DANE: {INST_DANE} | NIT: {INST_NIT}", f"Dirección: {INST_DIRECCION}", f"Fecha de impresión: {ahora().strftime('%Y-%m-%d %H:%M:%S')}"]


def estilo_excel(ws, header_row):
    amarillo = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid"); borde = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True); cell.border = borde
    for cell in ws[header_row]: cell.fill = amarillo; cell.font = Font(bold=True)
    for col, width in {"A":16,"B":32,"C":12,"D":26,"E":16,"F":14,"G":22}.items(): ws.column_dimensions[col].width = width


@app.route("/exportar_estudiantes")
def exportar_estudiantes():
    wb = Workbook(); ws = wb.active
    for l in encabezado(): ws.append([l])
    ws.append([]); header_row = ws.max_row + 1; ws.append(["CÓDIGO", "NOMBRE Y APELLIDO", "GRADO", "DIRECTOR DE GRUPO"])
    for e in Estudiante.query.order_by(Estudiante.grado.asc(), Estudiante.nombre.asc()).all(): ws.append([e.codigo, f"{e.nombre} {e.apellido}", e.grado, e.director])
    estilo_excel(ws, header_row); b = BytesIO(); wb.save(b); b.seek(0); return send_file(b, as_attachment=True, download_name="estudiantes_edutrack.xlsx")


@app.route("/exportar_excel_reportes")
def exportar_excel():
    wb = Workbook(); ws = wb.active
    for l in encabezado(): ws.append([l])
    ws.append([]); header_row = ws.max_row + 1; ws.append(["CÓDIGO", "NOMBRE Y APELLIDO", "GRADO", "DIRECTOR DE GRUPO", "FECHA", "HORA", "QUIÉN REPORTA"])
    for r in datos_reporte(): ws.append(r)
    estilo_excel(ws, header_row); b = BytesIO(); wb.save(b); b.seek(0); return send_file(b, as_attachment=True, download_name="reporte_ingresos_edutrack.xlsx")


@app.route("/exportar_pdf")
def exportar_pdf():
    b = BytesIO(); pdf = canvas.Canvas(b, pagesize=landscape(letter)); width, height = landscape(letter); y = height - 45
    try: pdf.drawImage("static/img/logo-colegio.png", 70, y - 65, width=65, height=65, preserveAspectRatio=True)
    except Exception: pass
    pdf.setFont("Helvetica-Bold", 14); pdf.drawCentredString(width / 2, y, INST_NOMBRE); y -= 17; pdf.setFont("Helvetica", 9)
    for linea in encabezado()[1:]: pdf.drawCentredString(width / 2, y, linea); y -= 13
    y -= 20; pdf.setFont("Helvetica-Bold", 12); pdf.drawString(40, y, "REPORTE DE INGRESOS"); y -= 22
    data = [["CÓDIGO", "NOMBRE Y APELLIDO", "GRADO", "DIRECTOR DE GRUPO", "FECHA", "HORA", "QUIÉN REPORTA"]] + datos_reporte()
    table = Table(data, colWidths=[70, 160, 60, 140, 80, 70, 120]); table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.yellow),("GRID", (0, 0), (-1, -1), 0.7, colors.black),("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),("ALIGN", (0, 0), (-1, -1), "CENTER"),("FONTSIZE", (0, 0), (-1, -1), 8)])); table.wrapOn(pdf, width, height); table.drawOn(pdf, 40, max(80, y - (len(data) * 22)))
    pdf.setFont("Helvetica", 8); pdf.drawCentredString(width / 2, 25, f"{APP_NAME} © 2026 | {SLOGAN} | Desarrollado por {DESARROLLADOR}"); pdf.save(); b.seek(0); return send_file(b, as_attachment=True, download_name="reporte_ingresos_edutrack.pdf")


@app.route("/exportar_word")
def exportar_word():
    doc = Document(); doc.add_paragraph("\n".join(encabezado())); doc.add_heading("REPORTE DE INGRESOS", level=1); t = doc.add_table(rows=1, cols=7); t.style = "Table Grid"; headers = ["CÓDIGO", "NOMBRE Y APELLIDO", "GRADO", "DIRECTOR DE GRUPO", "FECHA", "HORA", "QUIÉN REPORTA"]
    for i, h in enumerate(headers): t.rows[0].cells[i].text = h
    for r in datos_reporte():
        cells = t.add_row().cells
        for i, v in enumerate(r): cells[i].text = str(v)
    doc.add_paragraph(f"{APP_NAME} © 2026 | {SLOGAN}"); doc.add_paragraph(f"Desarrollado por {DESARROLLADOR}"); b = BytesIO(); doc.save(b); b.seek(0); return send_file(b, as_attachment=True, download_name="reporte_ingresos_edutrack.docx")


@app.route("/legal")
def legal(): return page("Legal", f"<div class='center'><section class='card'><h1>Aviso legal</h1><p>{APP_NAME} desarrollado por {DESARROLLADOR}.</p><a href='/login'>Volver</a></section></div>")


@app.route("/cookies")
def cookies(): return page("Cookies", "<div class='center'><section class='card'><h1>Cookies</h1><p>Se usan cookies técnicas de sesión.</p><a href='/login'>Volver</a></section></div>")


@app.route("/logout")
def logout(): session.clear(); return redirect("/login")


if __name__ == "__main__":
    with app.app_context(): inicializar_bd()
    app.run(debug=True, host="0.0.0.0")
